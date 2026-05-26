# -*- coding: utf-8 -*-
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import generate_practice_docx as g


OUTPUT_PATH = r"E:\Codex\webdev-exam\Отчет_по_практике_ПМ02_по_образцу.docx"


def add_assignment_page(document):
    for line in [
        "СОГЛАСОВАНО:",
        "Председатель ПЦК",
        "________/_______________",
        "«_____» _________ 20__ г.",
    ]:
        g.add_text(document, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    for _ in range(2):
        g.add_text(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    g.add_text(
        document,
        "ЗАДАНИЕ\nна производственную практику (по профилю специальности)\nпо профессиональному модулю ПМ.02 «Осуществление интеграции программных модулей»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        first_line=False,
    )
    g.add_text(
        document,
        "Специальность 09.02.07 «Информационные системы и программирование»",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    g.add_text(
        document,
        "Студент(ка) IV курса 22290907/1097 группы\nЖук Алексей Александрович",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    g.add_text(
        document,
        "Место прохождения практики: ________________________________________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    g.add_text(
        document,
        "Период прохождения практики\nс «12» января 2026 г.   по  «13» февраля 2026 г.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    g.add_text(
        document,
        "Виды работ, обязательные для выполнения:\nЗнакомство с предприятием\nУчастие в выработке требований к программному обеспечению\nУчастие в проектировании программного обеспечения с использованием специализированных программных пакетов\nУчастие в интеграции программных модулей\nРазработка рабочего проекта и технологической документации\nПриемо-сдаточные мероприятия",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    g.add_text(
        document,
        "Индивидуальное задание:\nВыполнить анализ исходного проекта веб-приложения интернет-магазина «Алтайская лавка», описать архитектуру, структуру базы данных, основные программные модули, алгоритмы, пользовательский интерфейс, механизмы интеграции и тестирования. Подготовить отчетную документацию, дневник практики, характеристику и вывод по результатам практики.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    g.add_text(
        document,
        "Задание выдал   «12» января 2026 г. ____________________\nЗадание получил «12» января 2026 г. ____________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )


def intro_in_sample_style():
    return [
        "В рамках производственной практики была выполнена работа по анализу и документированию программного продукта, предназначенного для автоматизации процессов интернет-магазина «Алтайская лавка». Система ориентирована на публикацию каталога товаров, обработку заказов, работу с корпоративными заявками, администрирование ассортимента и подготовку интеграции с внешними сервисами.",
        "На практике часть задач интернет-магазина могла бы вестись разрозненно: каталог товаров хранится отдельно от заказов, учет остатков ведется в сторонних таблицах, а взаимодействие с корпоративными клиентами фиксируется вручную. Такой подход затрудняет сопровождение ассортимента, повышает вероятность ошибок и усложняет интеграцию модулей. Рассматриваемый проект решает данные проблемы за счет единой архитектуры приложения, общей базы данных и выделенного сервисного слоя.",
        "Цель производственной практики - исследовать архитектуру и реализацию веб-приложения на Python (Flask), предназначенного для автоматизации работы интернет-магазина, и подготовить комплект технической документации по результатам анализа проекта.",
        "Для достижения цели были поставлены задачи: изучить предметную область интернет-магазина и назначение программного продукта; определить функциональные требования и основные пользовательские сценарии; проанализировать структуру репозитория, архитектуру приложения и состав программных модулей; исследовать используемые модели данных, алгоритмы и механизмы интеграции; рассмотреть проектирование пользовательского интерфейса, тестирование и документацию; оформить результаты работы в виде отчета по практике.",
    ]


def build_restyled_doc():
    document = Document()
    g.set_default_style(document)
    document.sections[0].different_first_page_header_footer = True

    g.title_page(document, diary=False)
    document.add_page_break()
    add_assignment_page(document)
    document.add_page_break()
    g.title_page(document, diary=True)
    document.add_page_break()
    g.build_diary_table(document)
    document.add_page_break()

    g.add_heading(document, "СОДЕРЖАНИЕ")
    toc_paragraph = document.add_paragraph()
    g.set_paragraph_format(toc_paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    g.add_toc(toc_paragraph)

    sections = deepcopy(g.REPORT_SECTIONS)
    sections[0] = (sections[0][0], intro_in_sample_style() + sections[0][1])

    for title, paragraphs in sections:
        document.add_page_break()
        g.add_heading(document, title)
        for paragraph in paragraphs:
            g.add_text(document, paragraph)

    document.add_page_break()
    g.add_heading(document, "ХАРАКТЕРИСТИКА СТУДЕНТА")
    for paragraph in g.CHARACTERISTIC:
        g.add_text(document, paragraph)

    document.add_page_break()
    g.add_heading(document, "ВЫВОД ПО ПРАКТИКЕ")
    for paragraph in g.PRACTICE_CONCLUSION:
        g.add_text(document, paragraph)

    document.add_page_break()
    g.add_heading(document, "СПИСОК ИСПОЛЬЗОВАННЫХ ТЕХНОЛОГИЙ")
    g.build_numbered_list(document, g.TECHNOLOGIES)

    document.add_page_break()
    g.add_heading(document, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    g.build_numbered_list(document, g.LITERATURE)

    for section in document.sections:
        section.footer.is_linked_to_previous = False
        footer_paragraph = section.footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        g.add_page_number(footer_paragraph)

    document.sections[0].different_first_page_header_footer = True
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_restyled_doc()
